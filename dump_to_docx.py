# dump_to_docx.py — only src/, prisma/, uploads/, middleware.ts, .env
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from pathlib import Path
import fnmatch

ROOT = Path(".").resolve()
OUT  = ROOT / "project_dump.docx"

# What to include
TARGET_SUBTREES = {"src", "prisma", "uploads"}
TARGET_FILES    = {"middleware.ts", ".env"}

# Skip these directories anywhere in the path
EXCLUDE_PARTS = {
    "node_modules",".git",".next","dist","build","out","coverage",
    ".turbo",".vercel",".vscode","__pycache__","venv",".venv"
}

# Text-ish file patterns to include inside the target folders
INCLUDE_GLOBS = [
    "*.ts","*.tsx","*.js","*.jsx","*.mjs","*.cjs",
    "*.css","*.scss","*.sass",
    "*.html","*.json","*.md","*.txt",
    "*.yml","*.yaml",
    "*.sql","*.prisma",
    "*.env","*.env.*"
]

def is_under_target_subtree(p: Path) -> bool:
    """True if the file lives under one of the TARGET_SUBTREES at repo root."""
    try:
        rel = p.relative_to(ROOT)
    except ValueError:
        return False
    return len(rel.parts) > 0 and rel.parts[0] in TARGET_SUBTREES

def is_target_root_file(p: Path) -> bool:
    """True if it's one of the named files at the repo root."""
    return p.parent == ROOT and p.name in TARGET_FILES

def should_exclude(p: Path) -> bool:
    return any(part in EXCLUDE_PARTS for part in p.parts)

def matches_globs(name: str) -> bool:
    return any(fnmatch.fnmatch(name, g) for g in INCLUDE_GLOBS)

def mono_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    f = r.font
    f.name = "Consolas"
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    f.size = Pt(10)

def main():
    doc = Document()
    doc.add_heading(str(ROOT.name) + " — selected files", level=1)

    files = []

    # 1) Add specific root files if present
    for fname in TARGET_FILES:
        p = ROOT / fname
        if p.exists() and p.is_file():
            files.append(p)

    # 2) Add files from the target subtrees that match text patterns
    for subtree in TARGET_SUBTREES:
        base = ROOT / subtree
        if not base.exists():
            continue
        for p in base.rglob("*"):
            if not p.is_file():
                continue
            if should_exclude(p):
                continue
            # Always allow .env files; otherwise check text globs
            if p.name == ".env" or matches_globs(p.name):
                files.append(p)

    # Make paths unique and sorted
    files = sorted(set(files), key=lambda x: x.as_posix())

    for f in files:
        rel = f.relative_to(ROOT)
        doc.add_heading(str(rel), level=2)
        try:
            content = f.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            content = f"[[ Could not read file: {e} ]]"
        mono_paragraph(doc, content)

    doc.save(OUT)
    print(f"Wrote {OUT}")

if __name__ == "__main__":
    main()
